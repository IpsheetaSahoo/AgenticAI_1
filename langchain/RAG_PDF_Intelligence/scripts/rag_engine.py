# rag_engine.py

from opensearch_connector import client
from embedder import generate_embeddings
from retriever_pipeline import retrieve_top_k
from reranking import mmr_rerank
from langchain_core.prompts import PromptTemplate
from docx import Document

# OpenAI
from langchain_openai import ChatOpenAI
import os
from dotenv import load_dotenv

# ---------- CONFIG ----------
TOP_K_RETRIEVAL = 10
TOP_N_FINAL = 3

# Load environment variables
load_dotenv()
os.environ['OPENAI_API_KEY'] = os.getenv('OPENAI_API_KEY')

# ---------- INIT LLM (Gemini) ----------
llm = ChatOpenAI(model="gpt-4o-mini")

# ---------- PROMPT ----------
QA_PROMPT = PromptTemplate.from_template(
    """
    You are a helpful assistant answering user queries based on provided context.

    Context:
    {context}

    Question:
    {question}

    Answer:
    """
)

def run_rag_pipeline(query_text: str):
    # Step 1: Embed the query
    query_chunk = {"content": query_text, "metadata": {}}
    embedded_query = generate_embeddings([query_chunk])[0]["embedding"]

    # Step 2: Get best index
    best_index = "pdf_hnsw_index"
    print(f"📌 Using best index: {best_index}")

    # Step 3: Retrieve top-K docs
    all_retrieved = retrieve_top_k(embedded_query)
    retrieved_docs = all_retrieved[best_index]

    # Step 4: MMR Reranking
    final_docs = mmr_rerank(embedded_query, retrieved_docs, top_n=TOP_N_FINAL)

    # Step 5: Build context and run LLM
    context_text = "\n\n".join([doc["text"] for doc in final_docs])
    formatted_prompt = QA_PROMPT.format(context=context_text, question=query_text)

    # Step 6: Run LLM   
    answer = llm.invoke(formatted_prompt)

    # Extract only the content
    if hasattr(answer, "content"):
        answer_text = answer.content
    else:
        answer_text = str(answer)


    print("\n🧠 Final Answer:")
    print(answer_text)

    print("\n📚 Context used:")
    for i, doc in enumerate(final_docs, 1):
        print(f"{i}. Page: {doc['metadata']['page_number']} | Chunk: {doc['metadata']['chunk_id']}")
        print(f"   Score: {round(doc['score'], 4)}")
        print(f"   Text: {doc['text'][:100]}...\n")

    return query_text, answer_text, final_docs


    # Save to DOCX
def save_answer_to_docx(query, answer, context_docs, output_path="../outputs/rag_response_output.docx"):
    doc = Document()

    doc.add_heading("RAG Answer Report", level=1)
    doc.add_heading("🔍 Query:", level=2)
    doc.add_paragraph(query)

    doc.add_heading("🧠 LLM Answer:", level=2)
    doc.add_paragraph(str(answer))  # ensure string

    doc.add_heading("📚 Supporting Context:", level=2)
    for i, doc_item in enumerate(context_docs, 1):
        score = round(doc_item["score"], 4)
        page = doc_item["metadata"]["page_number"]
        chunk = doc_item["metadata"]["chunk_id"]
        text = doc_item["text"]

        doc.add_paragraph(f"{i}. Page: {page} | Chunk: {chunk} | Score: {score}", style="List Number")
        doc.add_paragraph(text[:500] + ("..." if len(text) > 500 else ""), style="Intense Quote")

    doc.save(output_path)
    print(f"\n📄 Document saved as: {output_path}")

if __name__ == "__main__":
    sample_query = "How does the brain process images and what are the key areas involved in visual perception?"
    query_text, answer, final_docs = run_rag_pipeline(sample_query)
    save_answer_to_docx(query_text, answer, final_docs)